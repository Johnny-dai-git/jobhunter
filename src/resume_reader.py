"""Resume reading: parse PDF/DOCX/MD/TXT into plain text."""
from __future__ import annotations

from pathlib import Path
from typing import Iterable


SUPPORTED_EXTS = {".pdf", ".docx", ".md", ".txt"}
PAUSED_SUBDIR = "_paused"


def find_resume(resume_dir: Path) -> Path:
    """Find the latest resume file in resume_dir."""
    files = list_resumes(resume_dir)
    if not files:
        raise FileNotFoundError(
            f"No resume files found in {resume_dir} (supported: {sorted(SUPPORTED_EXTS)})"
        )
    return files[0]


def list_resumes(resume_dir: Path) -> list[Path]:
    """List all "active" resume files in resume_dir (excluding paused), sorted by mtime descending.

    Ignore internal cache files starting with underscore + subdirectory (_paused/).
    """
    if not resume_dir.exists():
        return []
    files = [
        p for p in resume_dir.iterdir()
        if p.is_file()
        and p.suffix.lower() in SUPPORTED_EXTS
        and not p.name.startswith("_")
        and not p.name.startswith(".")
    ]
    files.sort(key=lambda p: -p.stat().st_mtime)
    return files


def list_paused_resumes(resume_dir: Path) -> list[Path]:
    """List resumes in resume_dir/_paused/, sorted by mtime descending."""
    paused_dir = resume_dir / PAUSED_SUBDIR
    if not paused_dir.exists():
        return []
    files = [
        p for p in paused_dir.iterdir()
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS
    ]
    files.sort(key=lambda p: -p.stat().st_mtime)
    return files


def pause_resume_file(resume_dir: Path, filename: str) -> Path:
    """Move a resume to _paused/ subdirectory. Return new path."""
    src = resume_dir / filename
    if not src.exists() or not src.is_file():
        raise FileNotFoundError(f"Cannot find {filename}")
    if src.suffix.lower() not in SUPPORTED_EXTS:
        raise ValueError(f"Unsupported format {src.suffix}")
    paused_dir = resume_dir / PAUSED_SUBDIR
    paused_dir.mkdir(exist_ok=True)
    dst = paused_dir / filename
    if dst.exists():
        dst.unlink()  # Overwrite
    src.rename(dst)
    return dst


def unpause_resume_file(resume_dir: Path, filename: str) -> Path:
    """Restore a resume from _paused/ back to main directory."""
    paused_dir = resume_dir / PAUSED_SUBDIR
    src = paused_dir / filename
    if not src.exists():
        raise FileNotFoundError(f"Cannot find paused {filename}")
    dst = resume_dir / filename
    if dst.exists():
        raise FileExistsError(f"Resume with same name {filename} already exists, delete or rename it first")
    src.rename(dst)
    return dst


def delete_paused_resume(resume_dir: Path, filename: str) -> None:
    """Delete a paused resume."""
    paused_dir = resume_dir / PAUSED_SUBDIR
    target = paused_dir / filename
    if target.exists():
        target.unlink()


def read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    chunks: list[str] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            chunks.append(f"--- Page {i + 1} ---\n{text.strip()}")
    return "\n\n".join(chunks)


def read_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []

    # Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())

    # Tables (common for resume formatting)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def read_resume(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return read_pdf(path)
    if ext == ".docx":
        return read_docx(path)
    if ext in {".md", ".txt"}:
        return read_text(path)
    raise ValueError(f"Unsupported resume format: {ext}")


def parse_and_cache(resume_dir: Path) -> tuple[Path, str]:
    """Read latest resume, cache text as .cache.txt, return (source file path, text)."""
    src = find_resume(resume_dir)
    text = read_resume(src)
    cache_path = resume_dir / "_parsed.txt"
    cache_path.write_text(text, encoding="utf-8")
    return src, text


def load_cached(resume_dir: Path) -> str:
    """Load cached resume text; if not exists, parse on the spot."""
    cache_path = resume_dir / "_parsed.txt"
    if cache_path.exists():
        return cache_path.read_text(encoding="utf-8")
    _, text = parse_and_cache(resume_dir)
    return text


def read_materials(materials_dir: Path) -> str:
    """Read all supported format files in materials directory, concatenate into text.

    Each file is marked with filename for easy source distinction by model.
    Return empty string if directory doesn't exist or is empty.
    """
    if not materials_dir.exists():
        return ""

    files = sorted(
        [
            p for p in materials_dir.iterdir()
            if p.is_file()
            and p.suffix.lower() in SUPPORTED_EXTS
            and not p.name.startswith("_")
            and not p.name.startswith(".")
        ],
        key=lambda p: p.name,
    )
    if not files:
        return ""

    parts: list[str] = []
    for f in files:
        try:
            text = read_resume(f).strip()
            if text:
                parts.append(f"### [{f.name}]\n{text}")
        except Exception as e:
            parts.append(f"### [{f.name}]\n(Read failed: {e})")

    return "\n\n".join(parts)
